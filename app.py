"""
app.py
Servidor web Flask para el generador de informes de tasacion.

Endpoints:
  GET  /            -> sirve la pagina web (index.html)
  POST /procesar    -> recibe plantilla + peritacion (texto o Excel) +
                       cotizacion (opcional), devuelve el informe completo
"""

import os
from functools import wraps
import hmac
from io import BytesIO
from flask import (Flask, request, send_file, render_template, jsonify,
                   session, redirect, url_for)

from procesador import procesar


app = Flask(__name__)

# Limite de tamano de subida: 10 MB
app.config["MAX_CONTENT_LENGTH"] = 10 * 1024 * 1024

# Clave para firmar las sesiones (cookies). Se lee del entorno; si no
# esta, se genera una aleatoria (las sesiones se invalidan al reiniciar).
app.secret_key = os.environ.get("SECRET_KEY") or os.urandom(32)

# Clave de acceso compartida del estudio (variable de entorno en Render).
CLAVE_ACCESO = os.environ.get("APP_PASSWORD", "").strip()


def login_requerido(f):
    """Decorador: exige sesion iniciada para acceder a la ruta."""
    @wraps(f)
    def wrapper(*args, **kwargs):
        if not session.get("autenticado"):
            return redirect(url_for("login"))
        return f(*args, **kwargs)
    return wrapper


@app.route("/login", methods=["GET", "POST"])
def login():
    """Pantalla de ingreso con clave compartida."""
    if request.method == "POST":
        ingresada = request.form.get("password", "")
        # Comparacion en tiempo constante para evitar fugas por timing.
        if CLAVE_ACCESO and hmac.compare_digest(ingresada, CLAVE_ACCESO):
            session["autenticado"] = True
            return redirect(url_for("index"))
        return render_template("login.html", error=True)
    if session.get("autenticado"):
        return redirect(url_for("index"))
    return render_template("login.html", error=False)


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))

# Carpeta donde viven las plantillas virgenes precargadas
PLANTILLAS_DIR = os.path.join(os.path.dirname(__file__), "plantillas")

# Plantillas disponibles: clave -> nombre de archivo en PLANTILLAS_DIR
PLANTILLAS = {
    "agrosalta": "Informe_Tasacion_Agrosalta.xlsx",
    "cooperacion": "Informe_Tasacion_Cooperacion.xlsx",
}


def _cargar_plantilla(clave):
    """Lee los bytes de una plantilla precargada. Devuelve None si no existe."""
    nombre = PLANTILLAS.get(clave)
    if not nombre:
        return None
    ruta = os.path.join(PLANTILLAS_DIR, nombre)
    if not os.path.exists(ruta):
        return None
    with open(ruta, "rb") as f:
        return f.read()


# ============================================================
#  LECTURA DE ARCHIVOS
# ============================================================

def _texto_de_archivo(archivo):
    """
    Extrae el texto de un archivo de cotizacion subido.
    Soporta: .txt, .xlsx/.xls (Excel) y .docx (Word).
    Devuelve el texto plano, o "" si no se pudo leer.
    """
    if not archivo or archivo.filename == "":
        return ""

    nombre = archivo.filename.lower()
    datos = archivo.read()

    # --- Texto plano ---
    if nombre.endswith(".txt"):
        try:
            return datos.decode("utf-8", errors="ignore")
        except Exception:
            return ""

    # --- Excel ---
    if nombre.endswith((".xlsx", ".xls")):
        return _texto_de_excel(datos, nombre)

    # --- Word ---
    if nombre.endswith(".docx"):
        try:
            from docx import Document
            doc = Document(BytesIO(datos))
            partes = [p.text for p in doc.paragraphs if p.text.strip()]
            for tabla in doc.tables:
                for fila in tabla.rows:
                    celdas = [c.text.strip() for c in fila.cells
                              if c.text.strip()]
                    if celdas:
                        partes.append("\t".join(celdas))
            return "\n".join(partes)
        except Exception:
            return ""

    return ""


def _texto_de_excel(datos, nombre):
    """
    Extrae el contenido de un Excel volcando CADA celda con su
    coordenada (ej: 'A55=Paragolpe', 'B55=X', 'D55=95000').

    Esto le permite a Claude entender la grilla de daños con sus
    sectores y las marcas 'X', sin importar como este estructurada.

    Soporta .xlsx/.xlsm (via openpyxl) y .xls viejo (via xlrd).
    La deteccion del formato se hace por los primeros bytes (firma del
    archivo), no por la extension, porque a veces vienen archivos .xls
    renombrados como .xlsx (y viceversa).
    """
    # Firmas: xlsx/zip = "PK", xls viejo (OLE2) = D0 CF 11 E0
    es_xls_viejo = datos.startswith(b"\xd0\xcf\x11\xe0")
    es_xlsx = datos.startswith(b"PK")

    lineas = []
    try:
        if es_xls_viejo:
            # --- Formato .xls viejo: usar xlrd ---
            import xlrd
            wb = xlrd.open_workbook(file_contents=datos)
            for sh in wb.sheets():
                lineas.append("--- HOJA: %s ---" % sh.name)
                for r in range(sh.nrows):
                    celdas = []
                    for c in range(sh.ncols):
                        v = sh.cell_value(r, c)
                        if v in ("", None):
                            continue
                        # Numeros enteros sin el .0
                        if isinstance(v, float) and v == int(v):
                            v = int(v)
                        coord = _coord(c, r + 1)
                        celdas.append("%s=%s" % (coord, v))
                    if celdas:
                        lineas.append(" | ".join(celdas))
        elif es_xlsx:
            # --- Formato .xlsx: usar openpyxl ---
            from openpyxl import load_workbook
            wb = load_workbook(BytesIO(datos), data_only=True)
            if not wb.sheetnames:
                # xlsx con cero hojas (suele pasar con archivos
                # exportados en formato "Strict OOXML" desde sistemas
                # legacy). Avisar al usuario en vez de devolver vacio.
                raise ValueError("xlsx_sin_hojas")
            for ws in wb.worksheets:
                lineas.append("--- HOJA: %s ---" % ws.title)
                for fila in ws.iter_rows():
                    celdas = []
                    for celda in fila:
                        if celda.value in (None, ""):
                            continue
                        celdas.append("%s=%s" % (celda.coordinate,
                                                 celda.value))
                    if celdas:
                        lineas.append(" | ".join(celdas))
        else:
            raise ValueError("formato_no_reconocido")
    except ValueError:
        raise
    except Exception:
        return ""

    return "\n".join(lineas)


def _coord(col_idx, fila):
    """Convierte indice de columna (0=A) + fila a coordenada tipo 'B55'."""
    letra = ""
    n = col_idx
    while True:
        letra = chr(65 + n % 26) + letra
        n = n // 26 - 1
        if n < 0:
            break
    return "%s%d" % (letra, fila)


@app.route("/")
@login_requerido
def index():
    """Sirve la pagina principal."""
    return render_template("index.html")


@app.route("/procesar", methods=["POST"])
@login_requerido
def procesar_informe():
    """
    Recibe:
      - plantilla    (archivo .xlsx, obligatorio)
      - texto        (texto libre del peritaje, opcional)
      - excel_perit  (archivo Excel del peritaje, opcional)
      - cotizacion   (texto de la cotizacion de repuestos, opcional)
      - cotiz_file   (archivo de cotizacion .txt/.xlsx/.docx, opcional)
    Hace falta texto O excel_perit (al menos uno).
    Devuelve el archivo .xlsx completado para descargar.
    """
    try:
        # --- Plantilla: precargada (por clave) o subida manualmente ---
        plantilla_bytes = None

        clave = request.form.get("plantilla_clave", "").strip().lower()
        if clave:
            plantilla_bytes = _cargar_plantilla(clave)
            if plantilla_bytes is None:
                return jsonify({"error": "Plantilla no reconocida"}), 400
        elif "plantilla" in request.files and \
                request.files["plantilla"].filename != "":
            plantilla_file = request.files["plantilla"]
            if not plantilla_file.filename.lower().endswith(".xlsx"):
                return jsonify({"error": "La plantilla debe ser un archivo .xlsx"}), 400
            plantilla_bytes = plantilla_file.read()
        else:
            return jsonify({"error": "Elegi una plantilla (Agrosalta o Cooperacion)"}), 400

        # --- Peritaje: texto pegado y/o archivo Excel ---
        texto = request.form.get("texto", "").strip()

        if "excel_perit" in request.files:
            f = request.files["excel_perit"]
            if f and f.filename != "":
                try:
                    texto_excel = _texto_de_archivo(f)
                except ValueError as e:
                    motivo = str(e)
                    if motivo == "xlsx_sin_hojas":
                        msg = ("El archivo Excel está guardado en un formato "
                               "no compatible (suele pasar con archivos "
                               "exportados desde sistemas de aseguradoras). "
                               "Abrilo en Excel y usá 'Guardar como' → "
                               "'Libro de Excel (.xlsx)' o '.xls', y volvé "
                               "a subirlo.")
                    elif motivo == "formato_no_reconocido":
                        msg = ("No pude reconocer el formato del archivo "
                               "Excel. Asegurate de que sea un .xlsx o .xls "
                               "válido.")
                    else:
                        msg = "No se pudo leer el archivo Excel."
                    return jsonify({"error": msg}), 400
                if texto_excel:
                    # Se marca claramente que es una grilla de Excel,
                    # para que el procesador la interprete como tal.
                    bloque = ("\n\n=== PERITAJE DESDE EXCEL (GRILLA) ===\n"
                              + texto_excel)
                    texto = (texto + bloque).strip()

        if not texto:
            return jsonify({
                "error": "Carga el texto del peritaje o subi el Excel de peritacion"
            }), 400

        # --- Cotizacion: texto pegado y/o archivo ---
        cotizacion = request.form.get("cotizacion", "").strip()

        if "cotiz_file" in request.files:
            texto_archivo = _texto_de_archivo(request.files["cotiz_file"])
            if texto_archivo:
                cotizacion = (cotizacion + "\n\n" + texto_archivo).strip()

        # --- Armar el texto completo y procesar ---
        texto_completo = texto
        if cotizacion:
            texto_completo += (
                "\n\n=== COTIZACION DE REPUESTOS ===\n" + cotizacion)

        # --- Campos cargados a mano (pisan lo inferido) ---
        overrides = {}
        fecha_insp = request.form.get("fecha_inspeccion", "").strip()
        dir_insp = request.form.get("direccion_inspeccion", "").strip()
        loc_insp = request.form.get("localidad_inspeccion", "").strip()
        # El input date manda AAAA-MM-DD; el informe usa DD/MM/AAAA.
        if fecha_insp and len(fecha_insp) == 10 and fecha_insp[4] == "-":
            a, m, d = fecha_insp.split("-")
            fecha_insp = f"{d}/{m}/{a}"
        if fecha_insp:
            overrides["fechaInspeccion"] = fecha_insp
        if dir_insp:
            overrides["tallerDireccion"] = dir_insp
        if loc_insp:
            overrides["tallerLocalidad"] = loc_insp

        informe_bytes, data, _ = procesar(
            plantilla_bytes,
            texto_peritacion=texto_completo,
            overrides=overrides,
        )

        # --- Nombre del archivo de salida ---
        nro = data.get("numeroSiniestro") or data.get("dominio") or "sin-numero"
        etiqueta = {"agrosalta": "Agrosalta",
                    "cooperacion": "Cooperacion"}.get(clave, "Agrosalta")
        nombre = f"Informe Tasacion {etiqueta} - COMPLETO - {nro}.xlsx"

        return send_file(
            BytesIO(informe_bytes),
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            as_attachment=True,
            download_name=nombre,
        )

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"Error al procesar: {str(e)}"}), 500


@app.errorhandler(413)
def archivo_muy_grande(e):
    return jsonify({"error": "El archivo es demasiado grande (maximo 10 MB)"}), 413


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=True)
